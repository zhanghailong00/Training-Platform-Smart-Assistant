"""Word 文档生成组件：Markdown → Word 转换

支持两种方式：
1. pypandoc（优先，需要安装 pandoc）
2. python-docx 手动构建（兜底，不需要额外依赖）
"""

import os
import re
from pathlib import Path
from core.logger import get_logger

logger = get_logger(__name__)

# 下载目录
_DOWNLOAD_DIR = Path(__file__).resolve().parent.parent / "downloads"


def _ensure_download_dir():
    """确保下载目录存在"""
    _DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)


def markdown_to_docx(markdown_text: str, output_path: str = None) -> str:
    """
    将 Markdown 文本转换为 Word 文档。

    优先使用 pypandoc（质量好），如果不可用则使用 python-docx 手动构建。

    Args:
        markdown_text: Markdown 格式的文本
        output_path: 输出文件路径（可选，默认自动生成）

    Returns:
        Word 文件的路径
    """
    _ensure_download_dir()

    if output_path is None:
        import uuid
        output_path = str(_DOWNLOAD_DIR / f"{uuid.uuid4().hex}.docx")

    # 尝试使用 pypandoc
    try:
        return _convert_with_pypandoc(markdown_text, output_path)
    except Exception as e:
        logger.warning(f"pypandoc 转换失败，使用 python-docx 兜底: {e}")
        return _convert_with_python_docx(markdown_text, output_path)


def _convert_with_pypandoc(markdown_text: str, output_path: str) -> str:
    """使用 pypandoc 转换 Markdown 到 Word"""
    import pypandoc

    pypandoc.convert_text(
        markdown_text,
        'docx',
        format='md',
        outputfile=output_path,
        extra_args=['--from=markdown', '--to=docx'],
    )
    logger.info(f"Word 文件生成成功（pypandoc）: {output_path}")
    return output_path


def _add_rich_paragraph(doc, text: str, style_type: str = 'normal', level: int = 1):
    """
    添加一个段落，支持行内加粗 **text** 和行内代码 `code`。

    Args:
        doc: Document 对象
        text: 可能包含 **bold** 或 `code` 的文本
        style_type: 'normal' | 'heading' | 'bullet' | 'number'
        level: 标题级别（仅 heading 时有效）
    """
    from docx.shared import Pt, RGBColor

    # 根据样式类型创建段落
    if style_type == 'heading':
        p = doc.add_heading('', level=level)
    elif style_type == 'bullet':
        p = doc.add_paragraph('', style='List Bullet')
    elif style_type == 'number':
        p = doc.add_paragraph('', style='List Number')
    else:
        p = doc.add_paragraph()

    # 如果没有特殊格式，直接设置文本
    if '**' not in text and '`' not in text:
        p.add_run(text)
        return

    # 解析行内格式：交替匹配 **bold** 和 `code`
    pattern = r'(\*\*.*?\*\*|`[^`]+`)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        # 加粗 **text**
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(12)
        # 行内代码 `code`
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            p.add_run(part)


def _convert_with_python_docx(markdown_text: str, output_path: str) -> str:
    """使用 python-docx 手动构建 Word 文档（支持代码块、加粗等格式）"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 创建代码块样式（基于 Normal，等宽字体）
    code_style = doc.styles.add_style('CodeBlock', 1)  # 1 = paragraph
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(10)
    code_font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    code_paragraph_format = code_style.paragraph_format
    code_paragraph_format.space_before = Pt(2)
    code_paragraph_format.space_after = Pt(2)

    lines = markdown_text.split('\n')
    in_code_block = False
    code_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块开始/结束
        if stripped.startswith('```'):
            if in_code_block:
                # 结束代码块，写入所有收集的代码行
                for code_line in code_lines:
                    doc.add_paragraph(code_line, style='CodeBlock')
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        # 代码块内部
        if in_code_block:
            code_lines.append(line if line else ' ')
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 一级标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            _add_rich_paragraph(doc, stripped[2:], 'heading', 1)

        # 二级标题
        elif stripped.startswith('## '):
            _add_rich_paragraph(doc, stripped[3:], 'heading', 2)

        # 三级标题
        elif stripped.startswith('### '):
            _add_rich_paragraph(doc, stripped[4:], 'heading', 3)

        # 无序列表
        elif stripped.startswith('- ') or stripped.startswith('* '):
            _add_rich_paragraph(doc, stripped[2:], 'bullet')

        # 有序列表（带编号步骤）
        elif re.match(r'^\d+[\.\)]\s', stripped):
            text = re.sub(r'^\d+[\.\)]\s', '', stripped)
            _add_rich_paragraph(doc, text, 'number')

        # 普通段落（处理行内加粗、行内代码）
        else:
            _add_rich_paragraph(doc, stripped, 'normal')

        i += 1

    # 如果代码块没有闭合，也写入
    if code_lines:
        for code_line in code_lines:
            doc.add_paragraph(code_line, style='CodeBlock')

    doc.save(output_path)
    logger.info(f"Word 文件生成成功（python-docx）: {output_path}")
    return output_path


def cleanup_old_files(max_age_hours: int = 24):
    """清理过期的下载文件"""
    import time
    now = time.time()
    count = 0
    for f in _DOWNLOAD_DIR.iterdir():
        if f.is_file() and f.suffix == '.docx':
            age_hours = (now - f.stat().st_mtime) / 3600
            if age_hours > max_age_hours:
                f.unlink()
                count += 1
    if count:
        logger.info(f"清理了 {count} 个过期文件")